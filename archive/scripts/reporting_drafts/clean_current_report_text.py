#!/usr/bin/env python3
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[2]
INPUT = Path("/private/tmp/SIADS_696_Team_7_Final_Report_current_no_comments.docx")
OUTPUT = ROOT / "report" / "SIADS_696_Team_7_Final_Report_current_cleaned.docx"

REPLACEMENTS = {
    "Merge ACS data with CDC PLACES and AirData on county name [CHECK]": (
        "Merge CDC PLACES, AirData, ACS, NOAA, and Census geography using five-digit county GEOID as the primary key."
    ),
    "[CHECK - should we list ONLY the columns we used for supervised/unsupervised efforts?]": (
        "Table 1 lists the core columns used for the supervised and unsupervised modeling work; the full enriched schema is available in the processed dataset and output artifacts."
    ),
    "ACS [CHECK]": "ACS",
    (
        "Mental health needs exist at the intersection of societal and regulatory decision-making. "
        "On the one hand, being able to accurately predict future mental health needs can help individuals anticipate their [come back to this]. "
        "On a more macro-scale, this predictive ability can help governing figures with policy making and budget segmentation."
    ): (
        "Mental health needs exist at the intersection of community well-being and public resource planning. "
        "Improved predictive ability could help public health agencies identify where additional outreach, prevention, and behavioral-health capacity may be needed."
    ),
}


def set_paragraph_text(paragraph, text):
    style = paragraph.style
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)
    paragraph.style = style


def clean_paragraph(paragraph):
    text = paragraph.text
    new_text = text
    for old, new in REPLACEMENTS.items():
        new_text = new_text.replace(old, new)
    if new_text != text:
        set_paragraph_text(paragraph, new_text)


def main():
    doc = Document(INPUT)
    for paragraph in doc.paragraphs:
        clean_paragraph(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    clean_paragraph(paragraph)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
