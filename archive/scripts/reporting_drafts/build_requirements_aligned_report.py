#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT = Path("/Users/sophiaboettcher/Milestone-II-Project")
RUBRIC = PROJECT / "outputs" / "rubric_completion"
OUT = PROJECT / "SIADS 696 Team 7 Final Report - Requirements Draft.docx"


def fmt(x, digits=2):
    if pd.isna(x):
        return "--"
    if isinstance(x, (int, np.integer)):
        return f"{int(x):,}"
    return f"{float(x):,.{digits}f}"


def pct(x, digits=1):
    if pd.isna(x):
        return "--"
    return f"{float(x) * 100:.{digits}f}%"


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=60, start=100, bottom=60, end=100):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def make_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.8)
    sec.right_margin = Inches(0.8)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.0
    for name, size in [("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 11)]:
        st = styles[name]
        st.font.name = "Arial"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.paragraph_format.space_before = Pt(10)
        st.paragraph_format.space_after = Pt(4)
    if "Caption" in styles:
        cap = styles["Caption"]
        cap.font.name = "Arial"
        cap.font.size = Pt(9)
        cap.font.color.rgb = RGBColor(0, 0, 0)
        cap.paragraph_format.space_after = Pt(4)
    return doc


def para(doc, text):
    doc.add_paragraph(text)


def caption(doc, text):
    p = doc.add_paragraph(text, style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_table(doc, number, title, headers, rows, widths=None):
    caption(doc, f"Table {number}. {title}")
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = str(h)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    table_borders(table)
    for r, row in enumerate(table.rows):
        for c, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            if widths:
                set_cell_width(cell, widths[c])
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(8.5 if len(rows) > 8 else 9)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    if r == 0:
                        run.bold = True
    doc.add_paragraph("")


def add_fig(doc, number, path, title, width=6.2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    caption(doc, f"Figure {number}. {title}")


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def main():
    doc = make_doc()
    metrics = pd.read_csv(RUBRIC / "supervised_cv_metrics_with_sd.csv")
    importance = pd.read_csv(RUBRIC / "rf_feature_importance.csv")
    ablation = pd.read_csv(RUBRIC / "rf_feature_family_ablation.csv")
    sensitivity = pd.read_csv(RUBRIC / "rf_sensitivity_grid.csv")
    failures = pd.read_csv(RUBRIC / "supervised_failure_examples.csv")
    unsup = pd.read_csv(RUBRIC / "unsupervised_method_comparison.csv")
    anova = pd.read_csv(RUBRIC / "unsupervised_fmd_anova.csv")
    db_profile = pd.read_csv(RUBRIC / "dbscan_cluster_profile.csv")
    km_profile = pd.read_csv(RUBRIC / "kmeans_cluster_profile.csv")

    best = metrics.sort_values("holdout_rmse").iloc[0]
    baseline = metrics[metrics["model"].eq("Dummy mean baseline")].iloc[0]

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run("Frequent Mental Distress Prediction Using Air Pollution")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0, 0, 0)
    para(doc, "SIADS 696 Final Project Report | Team 7")
    p = doc.add_paragraph()
    r = p.add_run("Jaeah Kim, Kyle Rodriguez, Sophia Boettcher")
    r.bold = True

    doc.add_heading("Introduction", level=1)
    para(
        doc,
        "This project asks whether county-level environmental and socioeconomic profiles can predict future mental health need. "
        "The outcome is CDC PLACES frequent mental distress (FMD), the share of adults reporting poor mental health for 14 or more days in the past month. "
        "The practical motivation is resource planning: if environmental and social indicators can help identify future mental-health burden, counties and public-health agencies could better target outreach, prevention, and behavioral-health capacity."
    )
    para(
        doc,
        "We used both supervised and unsupervised learning. The supervised task predicts next-year FMD from current-year county features using Dummy, ElasticNet, Random Forest, and Gradient Boosting regressors. "
        "The unsupervised task compares density-based DBSCAN clustering with centroid-based K-Means clustering to identify socio-environmental county profiles. "
        "Compared with related studies that focus primarily on individual cohorts or single exposure families, our contribution is an enriched county-year panel and a leakage-safe predictive/exploratory pipeline combining pollution, climate, ACS socioeconomic variables, industry composition, and health-access proxies."
    )
    para(
        doc,
        f"The main supervised finding is cautious: the best model was {best['model']} with holdout RMSE {best['holdout_rmse']:.2f}, compared with the Dummy baseline RMSE {baseline['holdout_rmse']:.2f}. "
        "The improvement suggests weak signal, but the holdout R² remains near zero, so the model is not ready for policy forecasting. "
        "The main unsupervised finding is that both DBSCAN and K-Means separate a small set of structurally unusual counties from a dominant majority group; cluster-level FMD differences are exploratory rather than causal."
    )

    doc.add_heading("Related Work", level=1)
    add_table(
        doc,
        1,
        "Related studies and how this project differs",
        ["Study / resource", "1-2 sentence summary", "Difference from this project"],
        [
            [
                "Park et al., PM2.5/PM10 and depression risk in KLoSA",
                "This study found that prolonged exposure to particulate matter was associated with elevated depression risk among middle-aged and older adults in South Korea.",
                "It is individual/cohort-based and non-U.S.; our project is county-level, U.S.-wide, and predictive/exploratory rather than causal.",
            ],
            [
                "Lei et al., neighborhood environment and mental health in China",
                "This work links perceived pollution, neighborhood conditions, and socioeconomic status to mental health outcomes.",
                "Our project uses measured county-level AQI/climate variables and ML clustering rather than perception-based neighborhood measures.",
            ],
            [
                "Werder et al., Gulf Long-Term Follow-up Study",
                "This U.S.-based study examines environmental exposures such as PM2.5 and greenspace in relation to depression in the Southeastern U.S.",
                "Our project expands geographically across U.S. counties and adds supervised next-year prediction plus unsupervised county profiling.",
            ],
        ],
        [1900, 3700, 3760],
    )

    doc.add_heading("Data Sources", level=1)
    para(
        doc,
        "The final enriched panel contains 14,748 county-year rows from 2019-2023. Data were merged by five-digit county GEOID. "
        "The supervised target uses year t features to predict year t+1 FMD, and the unsupervised analysis uses the latest available year for cross-sectional county profiling."
    )
    add_table(
        doc,
        2,
        "Data sources, formats, variables, and coverage",
        ["Source", "Location / format", "Important variables", "Coverage"],
        [
            ["CDC PLACES", "data.cdc.gov CSV", "FMD prevalence, lower/upper CI", "County-level annual health estimates"],
            ["EPA AirData annual", "EPA CSV", "Days with AQI, Max AQI, Median AQI, pollutant-day variables", "Annual AQI summaries by county"],
            ["EPA daily AQI", "EPA ZIP/CSV", "Daily AQI mean/std/p90/max, threshold days above 100/150/200", "Daily monitor counties; sparse coverage"],
            ["ACS 5-year bulk tables", "Census pipe-delimited .dat", "Income, poverty, unemployment, education, insurance, industry composition", "County socioeconomic controls"],
            ["NOAA Climate at a Glance", "JSON", "Annual temperature, precipitation, anomalies", "County climate profiles"],
            ["Census boundaries", "Shapefile ZIP", "County geometries", "Mapping and spatial visualization"],
        ],
        [1700, 2100, 3600, 1960],
    )

    doc.add_heading("Feature Engineering", level=1)
    bullets(
        doc,
        [
            "Merged CDC PLACES, EPA, ACS, NOAA, and Census geography by county GEOID.",
            "Engineered next-year FMD by shifting mental_health_prevalence forward within each county.",
            "Excluded FMD and FMD confidence intervals from unsupervised clustering; FMD is reintroduced only as a held-out outcome.",
            "Median-imputed sparse EPA daily AQI features in modeling pipelines.",
            "Cleaned sentinel missing values in income-like fields before the rubric-completion rerun.",
            "Standardized numeric features for ElasticNet and unsupervised clustering.",
        ],
    )
    add_table(
        doc,
        3,
        "Feature families and leakage controls",
        ["Feature family", "Examples", "Modeling role"],
        [
            ["Identifiers", "geoid, county_name, state_name, year, geometry", "Merge keys, splits, mapping only; not predictors"],
            ["FMD outcome", "mental_health_prevalence, lower_ci, upper_ci, next_year_fmd", "Target/held-out outcome; excluded from clustering"],
            ["Annual AQI", "Max AQI, Median AQI, AQI category days, pollutant days", "Supervised and unsupervised predictors from current year"],
            ["Daily AQI", "daily_aqi_std, daily_aqi_p90, daily_aqi_days_over_100", "Extreme-event and variability predictors"],
            ["ACS socioeconomic", "median_income, poverty_count, unemployed_count, population", "Controls for community vulnerability"],
            ["Education/health/industry", "pct_uninsured, pct_bachelors_or_higher, pct_manufacturing", "Infrastructure and industrial-trait predictors"],
            ["Climate", "annual_avg_temp_f, annual_precip_inches, anomalies", "Environmental context beyond pollution"],
        ],
        [1900, 3900, 3560],
    )

    doc.add_heading("Part A. Supervised Learning", level=1)
    doc.add_heading("Methods", level=2)
    para(
        doc,
        "The supervised task is regression. We used a Dummy Regressor as a baseline, ElasticNet as a regularized linear model, Random Forest as a bagged tree ensemble, and Gradient Boosting as a sequential tree ensemble. "
        "These model families differ in assumptions and mechanisms: the Dummy model sets a minimum benchmark, ElasticNet tests linear signal under multicollinearity, Random Forest captures nonlinear interactions with lower tuning burden, and Gradient Boosting captures additive nonlinear structure."
    )
    para(
        doc,
        "Hyperparameter exploration used grouped cross-validation by feature year on the training period. The final holdout used 2022 features to predict 2023 FMD. "
        "Metrics are RMSE and MAE for error magnitude, plus R² for explained variance relative to a mean baseline."
    )
    add_table(
        doc,
        4,
        "Supervised model comparison with cross-validation mean and standard deviation",
        ["Model", "Holdout RMSE", "Holdout MAE", "Holdout R²", "CV RMSE mean", "CV RMSE SD"],
        [
            [r.model, fmt(r.holdout_rmse), fmt(r.holdout_mae), fmt(r.holdout_r2), fmt(r.cv_rmse_mean), fmt(r.cv_rmse_sd)]
            for r in metrics.itertuples(index=False)
        ],
        [2500, 1350, 1350, 1250, 1500, 1410],
    )

    doc.add_heading("Feature Importance and Ablation", level=2)
    para(
        doc,
        "For the best model, Random Forest, impurity-based feature importance ranked education and climate variables highly. "
        "A grouped ablation analysis retrained the model after removing each feature family. Removing education/health-access features increased holdout RMSE the most, suggesting that these variables contributed more consistently than annual or daily AQI features in the current panel."
    )
    add_fig(doc, 1, RUBRIC / "figure_rf_feature_importance.png", "Top Random Forest feature importances.", width=5.7)
    add_table(
        doc,
        5,
        "Random Forest grouped ablation",
        ["Removed feature family", "Features removed", "Holdout RMSE", "RMSE delta"],
        [[r.removed_family, int(r.features_removed), fmt(r.holdout_rmse), fmt(r.rmse_delta)] for r in ablation.itertuples(index=False)],
        [3100, 1700, 1700, 1700],
    )
    add_fig(doc, 2, RUBRIC / "figure_rf_ablation.png", "Family ablation for the Random Forest model.", width=5.7)

    doc.add_heading("Sensitivity, Tradeoffs, and Failure Analysis", level=2)
    para(
        doc,
        "Sensitivity analysis varied Random Forest minimum leaf size and max-feature sampling. The best setting used small leaves and sqrt feature sampling, while larger leaves generally reduced variance but worsened holdout RMSE. "
        "The tradeoff is interpretability and stability versus nonlinear flexibility: Random Forest performed best, but at the cost of reduced transparency relative to ElasticNet."
    )
    add_table(
        doc,
        6,
        "Random Forest sensitivity analysis, top settings",
        ["min_samples_leaf", "max_features", "Holdout RMSE", "Holdout R²"],
        [[int(r.min_samples_leaf), r.max_features, fmt(r.holdout_rmse), fmt(r.holdout_r2)] for r in sensitivity.head(6).itertuples(index=False)],
        [2100, 2100, 2400, 2400],
    )
    para(
        doc,
        "Failure analysis focused on the largest absolute residuals. The model mainly failed by underpredicting counties in the high-FMD tail, including Appalachian and rural counties where county-level pollution and ACS controls likely miss local behavioral-health infrastructure, disability burden, opioid-related distress, social isolation, and other contextual variables."
    )
    add_table(
        doc,
        7,
        "Specific high-error supervised examples",
        ["County", "State", "Actual FMD", "Prediction", "Residual", "Failure category"],
        [
            [r.county_name, r.state_name, fmt(r.next_year_fmd), fmt(r.prediction), fmt(r.residual), r.failure_category]
            for r in failures.head(6).itertuples(index=False)
        ],
        [1350, 1500, 1200, 1200, 1100, 3010],
    )

    doc.add_heading("Part B. Unsupervised Learning", level=1)
    doc.add_heading("Methods", level=2)
    para(
        doc,
        "The unsupervised analysis used two methods with different mechanisms: DBSCAN, a density-based clustering method that can label noise/outliers, and K-Means, a centroid-based partitioning method. "
        "Both used standardized non-FMD features; FMD was held out and compared after clusters were formed. PCA was used for visualization and cluster-search diagnostics."
    )
    add_table(
        doc,
        8,
        "Unsupervised method comparison",
        ["Method", "Selected parameters", "Clusters", "Noise %", "Silhouette embedding", "Silhouette scaled features"],
        [
            [r.method, r.params, int(r.clusters), fmt(r.noise_pct), fmt(r.silhouette_embedding), fmt(r.silhouette_scaled_features)]
            for r in unsup.itertuples(index=False)
        ],
        [1350, 2300, 1000, 1000, 1850, 1860],
    )
    add_table(
        doc,
        9,
        "Held-out FMD ANOVA across unsupervised clusters",
        ["Method", "F statistic", "p-value", "Groups"],
        [[r.method, fmt(r.f_stat), fmt(r.p_value, 4), int(r.groups)] for r in anova.itertuples(index=False)],
        [2000, 2200, 2200, 1800],
    )
    add_fig(doc, 3, RUBRIC / "figure_dbscan_pca_scatter.png", "DBSCAN clusters on the PCA embedding.", width=5.5)
    add_fig(doc, 4, RUBRIC / "figure_dbscan_cluster_profile.png", "DBSCAN standardized cluster profiles.", width=5.5)
    add_fig(doc, 5, RUBRIC / "figure_kmeans_pca_scatter.png", "K-Means clusters on the PCA embedding.", width=5.5)
    add_fig(doc, 6, RUBRIC / "figure_kmeans_cluster_profile.png", "K-Means standardized cluster profiles.", width=5.5)
    para(
        doc,
        "DBSCAN found a dominant county group plus outliers and a very small second cluster. K-Means produced two clusters with stronger silhouette scores, but one cluster was still much smaller than the other. "
        "The held-out FMD comparison was stronger for K-Means in this rerun, while DBSCAN was more useful for identifying unusual counties. These results are exploratory and do not imply causality."
    )

    doc.add_heading("Discussion", level=1)
    para(
        doc,
        "Part A showed that the enriched feature set contains some predictive signal, but not enough for confident next-year FMD forecasting. The best model improved over the Dummy baseline and approached zero R², but large residuals remained in high-distress counties. "
        "This was surprising because adding ACS, climate, and daily AQI features helped performance but did not solve the high-FMD tail."
    )
    para(
        doc,
        "Part B showed that the county feature space is dominated by one large group with smaller outlier groups. The main challenge was that density-based clustering did not produce balanced archetypes. "
        "Comparing DBSCAN with K-Means helped clarify the tradeoff between outlier detection and partitioning all counties into interpretable groups."
    )
    para(
        doc,
        "With more time, the most important extensions would be provider-density and healthcare-access data, rurality/urbanicity classifications, opioid burden, disability prevalence, greenspace, disaster exposure, and finer monthly or daily exposure windows."
    )

    doc.add_heading("Ethical Considerations", level=1)
    bullets(
        doc,
        [
            "Predictions are county-level planning signals, not individual diagnoses or risk scores.",
            "High-FMD counties should not be stigmatized; outputs should guide support, outreach, and further investigation.",
            "Sparse monitor coverage and modeled PLACES estimates should be communicated as uncertainty.",
            "Cluster labels should not be used to rank communities or justify resource withdrawal.",
            "Because the project is predictive and exploratory, it should not be interpreted as causal evidence that pollution alone causes FMD.",
        ],
    )

    doc.add_heading("Statement of Work", level=1)
    add_table(
        doc,
        10,
        "Team member contributions",
        ["Team member", "Contributions"],
        [
            ["Jaeah Kim", "Unsupervised learning analysis, clustering workflow, preprocessing, interpretation support."],
            ["Kyle Rodriguez", "Visualization support, model evaluation, residual and cluster interpretation, report synthesis."],
            ["Sophia Boettcher", "Dataset integration, imputation workflow, supervised baselines, external enrichment, final report drafting."],
            ["All team members", "EDA, data cleaning, model review, final interpretation, and presentation/report synthesis."],
        ],
        [2200, 7160],
    )

    doc.add_heading("References", level=1)
    refs = [
        "CDC. PLACES: Local Data for Better Health. https://www.cdc.gov/places/",
        "EPA. AirData Download Files. https://aqs.epa.gov/aqsweb/airdata/download_files.html",
        "U.S. Census Bureau. ACS Summary File. https://www2.census.gov/programs-surveys/acs/summary_file/",
        "U.S. Census Bureau. Cartographic Boundary Files. https://www.census.gov/geographies/mapping-files/time-series/geo/carto-boundary-file.html",
        "NOAA National Centers for Environmental Information. Climate at a Glance. https://www.ncei.noaa.gov/access/monitoring/climate-at-a-glance",
        "Park et al. Particulate matters (PM2.5, PM10) and depression risk among middle-aged and older adults using KLoSA, 2016-2020.",
        "Lei et al. The impact of neighborhood environment on mental health: Evidence from China.",
        "Werder et al. Gulf Long-Term Follow-up Study research on PM2.5, greenspace, and depression in the Southeastern United States.",
    ]
    for ref in refs:
        doc.add_paragraph(ref)

    doc.add_heading("Project Submission Links", level=1)
    para(doc, "GitHub repository: [insert repository URL before submission]")
    para(doc, "Google Docs version with comments enabled: [insert Google Docs URL before submission]")

    doc.add_heading("Appendix A. Final Feature List", level=1)
    feature_rows = []
    for r in importance.sort_values(["family", "feature"]).itertuples(index=False):
        feature_rows.append([r.feature, r.family, "Predictor"])
    feature_rows = [["next_year_fmd", "FMD outcome", "Supervised target"], ["mental_health_prevalence", "FMD outcome", "Held-out cluster evaluation"]] + feature_rows
    add_table(
        doc,
        11,
        "Final modeling feature schema",
        ["Column", "Feature family", "Role"],
        feature_rows,
        [3600, 3000, 2760],
    )

    doc.add_heading("Appendix B. Artifact Catalog", level=1)
    add_table(
        doc,
        12,
        "Project artifacts",
        ["Artifact", "Purpose"],
        [
            ["Supervised_Unsupervised_FMD_Analysis.ipynb", "Main modeling notebook"],
            ["county_panel_enriched.csv", "Final county-year modeling panel"],
            ["generate_rubric_analyses.py", "Supplemental analyses for rubric completion"],
            ["outputs/rubric_completion/*.csv", "CV, feature importance, ablation, sensitivity, failure, and unsupervised comparison tables"],
            ["outputs/rubric_completion/*.png", "Figures used in this report"],
        ],
        [3700, 5660],
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
